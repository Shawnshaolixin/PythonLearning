"""一次性资源脚本：md 面试文档 → PDF + Word fixture 生成 + 解析器往返验证。

C# 对照主线：
  本脚本 ≈ download_models.py 同款「运行一次的初始化脚本」（C#: 资源准备工具类）：
    把 .md 源文档转成 .pdf / .docx 两种格式，供 rag-ingest / rag-chunk-compare
    练习 pdfplumber 与 python-docx 解析 —— 真实文档只解析真实文件才有教学意义。

中文字体问题（本周最大技术坑，README 有完整记录）：
  reportlab 的 STSong-Light 是 CID 字体，**不内嵌字形**，pdfplumber 在缺该字体
  的机器上提取会失败/为空（pdfplumber issue #508）—— 教学演示直接崩。
  选 fpdf2：显式注册 TTF 后，fpdf2 用 fontTools 做字形子集化并内嵌到 PDF
  （带 ToUnicode CMap），pdfplumber 提取 100% 可靠。
  字体来源：优先 Windows 系统 SimHei（零下载），失败提示放 Noto / 设环境变量。

输出：src/rag_core/data/samples/*.pdf + *.docx（data/ 已 gitignore，可随时重建）
运行：uv run python -m rag_core.make_samples
"""

from __future__ import annotations  # C#: （等价）—— 让类型注解延迟求值

import os  # C#: System.Environment（环境变量）
import re  # C#: System.Text.RegularExpressions
import sys  # C#: System.Environment.Exit
from pathlib import Path  # C#: System.IO

from docx import Document  # C#: Open XML SDK —— 生成 Word 文档
from docx.shared import Pt  # C#: 字号（点）设置
from fpdf import FPDF  # C#: PdfSharp —— 画布式 PDF 生成

from rag_core.golden_qa import GOLDEN_QA, normalize  # C#: 复用静态数据 + 工具方法
from rag_core.parsers import parse_file  # C#: 复用解析分发器

DATA_DIR = Path(__file__).parent / "data"  # C#: 数据目录（已 gitignore）
OUT_DIR = DATA_DIR / "samples"  # C#: fixture 输出目录
FONTS_DIR = DATA_DIR / "fonts"  # C#: 本地字体备份目录（可手动拷入）

# 源文档：两份真实面试文档（已提交仓库）
SOURCES: list[Path] = [
    Path(__file__).parent / "game_interview_guide.md",
    Path(__file__).parent / "game_interview_100_questions.md",
]

# 中文字体候选列表（按优先级）：环境变量 → 本地备份 → Windows 系统字体
FONT_CANDIDATES: list[str] = [
    os.environ.get("RAG_CJK_FONT", ""),  # C#: 环境变量覆盖（部署时指定）
    str(FONTS_DIR / "simhei.ttf"),
    r"C:\Windows\Fonts\simhei.ttf",  # C#: 系统字体（本机 win32）
    r"C:\Windows\Fonts\msyh.ttc",  # 微软雅黑（TTC，fpdf2 2.7.6+ 支持）
]

HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")  # C#: ^(#{1,6})\s+ 前缀
NUMBERED_RE = re.compile(r"^\d+[.、]\s*(.*)$")  # C#: 编号列表项前缀
BULLET_RE = re.compile(r"^[-*+]\s+(.*)$")  # C#: 无序列表项前缀
QUOTE_RE = re.compile(r"^>\s*(.*)$")  # C#: 引用前缀
TABLE_ROW_RE = re.compile(r"^\s*\|.*\|.*$")  # C#: 表格行（PDF 提取表格不可靠，fixture 直接丢弃）
HR_RE = re.compile(r"^\s*(-{3,}|\*{3,})$")  # C#: 水平分割线
LINK_RE = re.compile(r"\[([^\]]+)\]\([^)]*\)")  # C#: markdown 链接 → 只留显示文本


def strip_line(raw: str) -> tuple[str, str] | None:
    """剥离一行的 markdown 标记，返回 (纯文本, 行类型)。返回 None = 整行丢弃。

    C#: private (string, LineKind)? StripLine(string raw) —— 返回元组或 null
    教学点：**只删标记，绝不改动正文汉字** —— 金标短语必须逐字存活（verify_roundtrip 把关）。
    行类型："h1"|"h2"|"h3"|"list"|"para"|"empty" —— 供 docx 生成时映射样式。
    """
    if not raw.strip():
        return ("", "empty")
    # 表格行 / 水平线：fixture 直接丢弃（解析回来结构也保不住，诚实演示）
    if TABLE_ROW_RE.match(raw) or HR_RE.match(raw):
        return None

    text = raw  # C#: string text = raw;
    kind = "para"

    m = HEADING_RE.match(raw)  # C#: Regex.Match
    if m:
        level = min(len(m.group(1)), 3)  # C#: Math.Min(标题层数, 3) —— docx 只有 9 级样式
        text = m.group(2).strip()
        kind = f"h{level}"
    elif m := BULLET_RE.match(raw):  # C#: C# 8 的 pattern matching 也支持这种写法
        text = m.group(1).strip()
        kind = "list"
    elif m := NUMBERED_RE.match(raw):
        text = m.group(1).strip()
        kind = "list"
    elif m := QUOTE_RE.match(raw):
        text = m.group(1).strip()
        kind = "para"

    # 行内标记：加粗 **、行内代码 `、链接 [t](u) → 显示文本（C#: Regex.Replace 多次）
    text = re.sub(r"\*\*", "", text)
    text = text.replace("`", "")
    text = LINK_RE.sub(r"\1", text)
    # 脏数据清洗：源文档里混入的 U+FFFD 替换字符（编码损坏标记，不是内容）
    # 留着会让 fpdf2 报 missing glyph 且提取断字 —— 直接删（C#: 数据清洗层）
    text = text.replace("�", "")
    return (text.strip(), kind)


def md_to_plain(md_text: str) -> str:
    """md 全文 → 剥离标记后的纯文本（fixture 共用）。

    C#: string MdToPlain(string md) —— LINQ 过滤 + Select 映射
    """
    lines: list[str] = []
    for raw in md_text.splitlines():  # C#: foreach (var raw in md.Split('\n'))
        result = strip_line(raw)
        if result is not None:  # C#: if (result != null)
            text, _ = result
            lines.append(text)
    return "\n".join(lines)  # C#: string.Join("\n", lines)


def ensure_cjk_font(pdf: FPDF) -> str:
    """在候选列表里找到第一个可用的中文字体并注册到 pdf 实例。

    C#: FontResolver —— 逐个尝试 FontCollection 注册，第一个成功者生效（含降级策略）
    教学点：字体解析器是「候选 + 冒烟注册」模式，第一个失败的候选立即回退，
    全部失败时给出手动解决方案而不是静默生成乱码 PDF。
    """
    for path in FONT_CANDIDATES:
        if not path:  # C#: string.IsNullOrEmpty —— 环境变量未设置
            continue
        if not Path(path).exists():  # C#: File.Exists(path)
            continue
        try:
            pdf.add_font("CJK", "", path)  # C#: 冒烟注册：能解析字体文件即认为可用
            return path
        except Exception:  # C#: catch (Exception) —— TTC 取错 face 等都可能失败
            continue

    # C#: 全部失败 → 明确的错误提示（比静默乱码好一万倍）
    raise RuntimeError(
        "找不到可用的中文字体（生成 PDF 需要 TTF 字体）。请任选其一：\n"
        f"  1. 下载 Noto Sans SC TTF 放到 {FONTS_DIR}/\n"
        "  2. 设置环境变量 RAG_CJK_FONT=字体路径\n"
    )


def md_to_docx(md_path: Path, out_path: Path) -> None:
    """md → Word：按行类型映射到 Word 样式（标题/列表/正文）。

    C#: Open XML SDK 生成 Word 文档（ParagraphProperties + StyleId）
    """
    doc = Document()  # C#: 新建文档（默认模板自带 Heading/List 样式）
    # 中文字体设置：默认模板西文字体，显式设成宋体/黑体（C#: 文档默认样式）
    style = doc.styles["Normal"]
    style.font.name = "宋体"  # C#: 文档默认字体
    style.font.size = Pt(10.5)  # C#: 五号

    for raw in md_path.read_text(encoding="utf-8").splitlines():
        result = strip_line(raw)
        if result is None:
            continue
        text, kind = result
        if kind == "empty":
            continue
        if kind.startswith("h"):
            doc.add_heading(text, level=int(kind[1]))  # C#: Heading 1-3 样式
        elif kind == "list":
            doc.add_paragraph(text, style="List Bullet")  # C#: ListBullet 样式
        else:
            doc.add_paragraph(text)  # C#: 普通段落

    out_path.parent.mkdir(parents=True, exist_ok=True)  # C#: Directory.CreateDirectory
    doc.save(out_path)  # C#: doc.SaveAs(path)


def md_to_pdf(md_path: Path, out_path: Path) -> None:
    """md → PDF：每个非空行一个多行单元格段落（fpdf2 自动换行 + 自动分页）。

    C#: PdfSharp 画布式 —— 逐段 DrawString，换页判断自己写
    fpdf2 的优势：multi_cell 自动换行分页，且 TTF 字体子集化内嵌，提取可靠。
    """
    pdf = FPDF(format="A4")  # C#: PdfDocument(A4)
    font_path = ensure_cjk_font(pdf)  # C#: 注册字体（候选回退）
    pdf.set_margins(15, 15, 15)  # C#: 页边距 15mm
    pdf.add_page()  # C#: 第一页
    pdf.set_font("CJK", "", 10.5)  # C#: 五号字

    for raw in md_path.read_text(encoding="utf-8").splitlines():
        result = strip_line(raw)
        if result is None:
            continue
        text, kind = result
        if kind == "empty":
            pdf.ln(2)  # C#: 空行 → 段落间距（PDF 提取后段落边界会丢失，教学点）
            continue
        if kind.startswith("h"):
            size = {"h1": 16, "h2": 14, "h3": 12}[kind]  # C#: 标题分级字号
            pdf.set_font("CJK", "", size)
            pdf.multi_cell(w=0, h=7, text=text, new_x="LMARGIN", new_y="NEXT")  # C#: DrawString
            pdf.set_font("CJK", "", 10.5)  # 恢复正文字号
        else:
            pdf.multi_cell(w=0, h=5.5, text=text, new_x="LMARGIN", new_y="NEXT")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    pdf.output(str(out_path))  # C#: pdf.Save(path)


def verify_roundtrip() -> list[tuple[str, bool, bool]]:
    """解析器往返验证：解析刚生成的 PDF/Word，检查金标短语是否逐字存活。

    C#: 断言式验证 —— 解析产物必须能找回关键短语，否则退出非 0
    教学点：这就是「解析器验证」—— PDF 逐行提取（长句可能拆行、表格丢失），
    Word 段落级保存（通常 100% 存活）。命中判定统一空白归一化（normalize）。
    """
    pdf_texts: dict[str, str] = {}
    docx_texts: dict[str, str] = {}
    for src in SOURCES:
        stem = src.stem  # C#: Path.GetFileNameWithoutExtension
        # 走 rag_core.parsers.parse_file —— 与 rag-ingest 完全相同的解析链路！
        pdf_texts[stem] = parse_file(OUT_DIR / f"{stem}.pdf").text
        docx_texts[stem] = parse_file(OUT_DIR / f"{stem}.docx").text

    results: list[tuple[str, bool, bool]] = []
    for qa in GOLDEN_QA:
        stem = qa.source.removesuffix(".md")  # C#: Replace(".md", "")
        ans = normalize(qa.answer)  # C#: 空白归一化后判定
        pdf_ok = ans in normalize(pdf_texts[stem])
        docx_ok = ans in normalize(docx_texts[stem])
        results.append((qa.answer, pdf_ok, docx_ok))
    return results


def main() -> None:
    """生成 4 个 fixture（幂等）+ 往返验证 + 打印结果。

    C#: Main() —— 资源初始化脚本入口
    """
    # 幂等检查：4 个输出都已有非空文件 → 跳过生成（C#: 同 download_models.ensure_model）
    outputs = [OUT_DIR / f"{src.stem}.pdf" for src in SOURCES] + [
        OUT_DIR / f"{src.stem}.docx" for src in SOURCES
    ]
    missing = [p for p in outputs if not (p.exists() and p.stat().st_size > 0)]
    if not missing:
        print("fixture 已存在，跳过生成（如需重建请删除 data/samples/）")
    else:
        print(f"生成 fixture → {OUT_DIR}/")
        for src in SOURCES:
            md_to_docx(src, OUT_DIR / f"{src.stem}.docx")
            md_to_pdf(src, OUT_DIR / f"{src.stem}.pdf")
            print(f"  {src.name}: docx + pdf 完成")

    # 注意：输出用 ASCII（OK/XX）而非 ✓/✗ —— Windows GBK 控制台打不出 Unicode 符号
    print("\n解析器往返验证（金标短语逐字存活检查）：")
    print(f"{'短语':<34}{'PDF':<6}{'Word':<6}")
    all_ok = True
    for phrase, pdf_ok, docx_ok in verify_roundtrip():
        print(f"{phrase[:32]:<34}{'OK' if pdf_ok else 'XX':<6}{'OK' if docx_ok else 'XX':<6}")
        all_ok = all_ok and pdf_ok and docx_ok

    if not all_ok:
        # C#: Environment.Exit(1) —— 验证失败必须让调用方看见
        print("\nXX 有短语在往返中丢失！请检查（可能需要从 golden_qa.BACKUP_ANSWERS 替换）。")
        sys.exit(1)
    print("\nOK 全部金标短语通过 PDF + Word 往返验证")


if __name__ == "__main__":  # C#: Main() 入口方法
    main()
