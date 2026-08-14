"""文档解析器 —— pdf / docx / md / txt → 统一文本结构。

C# 对照主线：
  本模块 ≈ C# 里的文档解析服务（策略模式 + 工厂分发）：
    ParsedDocument ParseFile(string path)
  三个解析器对应 C# 生态：
    .pdf  → iTextSharp.PdfTextExtractor / PdfSharp（页面流，逐页取文本）
    .docx → Open XML SDK（DocumentFormat.OpenXml.Wordprocessing，段落流）
    .md/.txt → File.ReadAllText（直读）

教学点（面试会问「RAG 的 Ingest 环节容易出什么问题」）：
  1. PDF 是「页面流」：按物理行/页提取，表格、多栏会错乱 —— 结构信息在提取时丢失
  2. Word 是「段落流」：段落级保存，结构信息比 PDF 完整
  3. 解析器职责单一：只做「文件 → 纯文本」，不做 markdown 剥离（那是 fixture
     生成的关注点，见 make_samples.py —— C# 分层思想）
"""

from dataclasses import dataclass  # C#: record（不可变数据载体）
from pathlib import Path  # C#: System.IO

import pdfplumber  # C#: PDF 文本提取（对标 iTextSharp）
from docx import Document  # C#: Open XML SDK —— 读 Word 文档


@dataclass
class ParsedDocument:
    """一次解析的结果。C#: public record ParsedDocument(string Source, string Format, string Title, string Text, int PageCount, int ParagraphCount);"""

    source: str  # 文件路径（C#: SourcePath）
    format: str  # "pdf" | "docx" | "md" | "txt"（C#: 枚举 or 字符串判别）
    title: str  # 标题：md 取第一个 "# " 标题，其余取文件名 stem（C#: FileNameWithoutExtension）
    text: str  # 解析出的纯文本（C#: FullText）
    page_count: int = 0  # pdf 页数；其他格式 0
    paragraph_count: int = 0  # docx 段落数 / md-txt 行数；pdf 0


def parse_file(path: str | Path) -> ParsedDocument:
    """按扩展名分发到对应解析器（唯一对外入口）。

    C#: public ParsedDocument ParseFile(string path)
        switch (Path.GetExtension(path).ToLower()) { case ".pdf": ... }
    """
    p = Path(path)
    suffix = p.suffix.lower()  # C#: Path.GetExtension(path).ToLower()
    if suffix == ".pdf":
        return parse_pdf(p)
    if suffix == ".docx":
        return parse_docx(p)
    if suffix in (".md", ".txt"):
        return parse_plain(p)
    # C#: throw new ArgumentException($"不支持的文档格式: {suffix}")
    raise ValueError(f"不支持的文档格式: {suffix}（支持 .pdf / .docx / .md / .txt）")


def parse_pdf(path: Path) -> ParsedDocument:
    """PDF 解析：逐页 extract_text() 拼接（页间插空行保留页边界）。

    C#: iTextSharp 逐页 PdfTextExtractor.GetTextFromPage(...) 拼接
    教学点：PDF 是页面流 —— 每物理行/物理页是独立文本块，
    长句子可能被拆行、表格单元格顺序错乱，这是 PDF 解析的固有代价。
    """
    pages: list[str] = []
    page_count = 0
    with pdfplumber.open(path) as pdf:  # C#: using (var pdf = PdfReader.Open(...))
        page_count = len(pdf.pages)
        for page in pdf.pages:  # C#: foreach (var page in pdf.Pages)
            text = page.extract_text() or ""  # C#: ?? ""（无文本页返回 None）
            pages.append(text)

    text = "\n\n".join(pages)  # C#: string.Join("\n\n", pages) —— 页间空行分隔
    return ParsedDocument(
        source=str(path),
        format="pdf",
        title=path.stem,  # C#: Path.GetFileNameWithoutExtension
        text=text,
        page_count=page_count,
    )


def parse_docx(path: Path) -> ParsedDocument:
    """Word 解析：遍历段落取文本，段落流（比 PDF 结构完整）。

    C#: Open XML SDK —— 遍历 WordprocessingDocument 的 Body.Paragraphs 取每个 Paragraph 的 InnerText
    教学点：Word 是段落流 —— 段落边界原样保留，不存在「拆行」问题。
    """
    doc = Document(path)  # C#: 打开 Word 文档（OpenXML 内存模型）
    lines = [p.text for p in doc.paragraphs if p.text.strip()]  # C#: doc.Paragraphs.Where(p => !string.IsNullOrWhiteSpace(p.Text)).Select(p => p.Text)

    text = "\n".join(lines)  # C#: string.Join("\n", lines)
    return ParsedDocument(
        source=str(path),
        format="docx",
        title=path.stem,
        text=text,
        paragraph_count=len(lines),
    )


def parse_plain(path: Path) -> ParsedDocument:
    """md / txt 直读：纯文本读入，不做 markdown 剥离（职责分离）。

    C#: File.ReadAllText(path, Encoding.UTF8)
    教学点：md 的 markdown 标记（#、**、列表）留给切片环节处理或 fixture
    生成时剥离 —— 解析器只负责「读文件」，单职责原则。
    """
    text = path.read_text(encoding="utf-8")  # C#: File.ReadAllText(path, Encoding.UTF8)

    # 标题：第一个 "# " 行（C#: text.Split('\n').FirstOrDefault(l => l.StartsWith("# "))）
    title = ""
    for line in text.splitlines():  # C#: foreach (var line in text.Split('\n'))
        if line.startswith("# "):  # C#: line.StartsWith("# ")
            title = line[2:].strip()  # C#: line.Substring(2).Trim()
            break
    if not title:
        title = path.stem  # C#: 回退：文件名

    return ParsedDocument(
        source=str(path),
        format=path.suffix.lower().lstrip("."),  # C#: "md" | "txt"
        title=title,
        text=text,
        paragraph_count=len([l for l in text.splitlines() if l.strip()]),
    )
