"""parsers 测试 —— Week 2 文档解析（真实文件系统 fixture，不需要 mock）。

C# 对照主线：
  tmp_path fixture                        ≈ xUnit 的临时目录（每个用例独立）
  生成微型 docx/pdf 再解析               ≈ 真实文件往返测试（OpenXML 生成 → 读取）
  无中文字体时 pytest.skip                ≈ 依赖缺失时的条件跳过（[Fact(Skip=...)]）

测试策略：不 mock —— 解析器就是「文件 → 文本」，喂真实的小文件最诚实。
PDF 用例依赖中文字体（ensure_cjk_font），本机没有字体则跳过而不是失败。
"""

from pathlib import Path  # C#: System.IO

import pytest  # C#: using Xunit;

from src.rag_core.golden_qa import normalize  # C#: 空白归一化（复用）
from src.rag_core.make_samples import ensure_cjk_font  # C#: 字体候选解析器（复用）
from src.rag_core.parsers import parse_file  # C#: 分发器（被测对象）


def test_parse_md_dispatch(tmp_path):
    """md 分发：format/标题（首个 # 行）/正文/段落数。"""
    md = tmp_path / "doc.md"
    md.write_text("# 测试标题\n\n正文内容第一行。\n正文内容第二行。", encoding="utf-8")

    parsed = parse_file(md)

    assert parsed.format == "md"
    assert parsed.title == "测试标题"
    assert "正文内容第一行。" in parsed.text
    assert parsed.paragraph_count == 3  # 非空行：标题 + 正文两行


def test_parse_txt_dispatch(tmp_path):
    """txt 分发：格式正确 + 文本原样。"""
    txt = tmp_path / "note.txt"
    txt.write_text("纯文本内容。", encoding="utf-8")

    parsed = parse_file(txt)

    assert parsed.format == "txt"
    assert parsed.text == "纯文本内容。"


def test_parse_unknown_extension_raises(tmp_path):
    """未知扩展名 → ValueError（C#: Assert.Throws<ArgumentException>）。"""
    bad = tmp_path / "doc.xyz"
    bad.write_text("随便什么", encoding="utf-8")

    with pytest.raises(ValueError):
        parse_file(bad)


def test_md_title_falls_back_to_stem(tmp_path):
    """没有 # 标题的 md → title 回退文件名（C#: Path.GetFileNameWithoutExtension）。"""
    md = tmp_path / "noheading.md"
    md.write_text("没有标题的文档。", encoding="utf-8")

    assert parse_file(md).title == "noheading"


def test_docx_roundtrip(tmp_path):
    """Word 往返：python-docx 生成 → parse_file 找回段落文本。

    C#: Open XML SDK 生成文档 → DocumentFormat.OpenXml 读取（真实往返）。
    """
    from docx import Document  # C#: 生成 Word 文档（写端）

    doc = Document()
    doc.add_paragraph("中文测试段落一。")
    doc.add_paragraph("中文测试段落二。")
    out = tmp_path / "sample.docx"
    doc.save(out)

    parsed = parse_file(out)

    assert parsed.format == "docx"
    assert parsed.paragraph_count == 2  # C#: 段落流 —— 段落数原样保留
    assert "中文测试段落一。" in parsed.text
    assert "中文测试段落二。" in parsed.text


def test_pdf_roundtrip(tmp_path):
    """PDF 往返：fpdf2 + 中文字体生成 → parse_file 提取（归一化后判定）。

    C#: PdfSharp 生成 → iTextSharp 提取（生成→提取闭环）。
    无中文字体时跳过（C#: 依赖缺失的条件跳过）—— 本地 Windows 有 SimHei，会真正执行。
    """
    from fpdf import FPDF  # C#: PdfSharp（写端）

    pdf = FPDF(format="A4")
    try:
        ensure_cjk_font(pdf)  # C#: 候选回退 —— 失败抛 RuntimeError
    except RuntimeError as e:  # C#: catch —— 无可用中文字体
        pytest.skip(f"无中文字体，跳过 PDF 往返测试：{e}")

    pdf.add_page()
    pdf.set_font("CJK", "", 12)
    pdf.multi_cell(w=0, h=8, text="中文 PDF 往返测试。你好世界！")
    out = tmp_path / "sample.pdf"
    pdf.output(str(out))

    parsed = parse_file(out)

    assert parsed.format == "pdf"
    assert parsed.page_count >= 1
    # PDF 是逐行提取：行尾换行可能拆句 → 归一化后判定（C#: 消除空白噪声）
    assert "中文PDF往返测试。你好世界！" in normalize(parsed.text)
